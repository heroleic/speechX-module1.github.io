from flask import Flask, render_template, request, jsonify, send_file
import requests
import os
from datetime import datetime
import uuid
from docx import Document
from docx.shared import Inches
import json
from dotenv import load_dotenv

# 加载环境变量
load_dotenv()

# 初始化Flask应用
app = Flask(__name__)

# DeepSeek API配置
DEEPSEEK_API_KEY = os.getenv('DEEPSEEK_API_KEY', 'your_api_key_here')  # 从环境变量获取API密钥，默认值为占位符
DEEPSEEK_API_URL = "https://api.deepseek.com/v1/chat/completions"  # DeepSeek API的URL

# 确保下载目录存在
os.makedirs('downloads', exist_ok=True)  # 创建用于存储下载文件的目录

def call_deepseek_api(prompt, temperature=0.3):
    """
    调用DeepSeek API进行文本处理
    :param prompt: 输入的提示文本
    :param temperature: 控制生成文本的随机性，默认为0.3
    :return: API返回的文本内容
    """
    headers = {
        "Authorization": f"Bearer {DEEPSEEK_API_KEY}",  # 设置API请求头
        "Content-Type": "application/json"
    }
    
    payload = {
        "model": "deepseek-chat",  # 指定模型
        "messages": [{"role": "user", "content": prompt}],  # 用户输入的消息
        "temperature": temperature,  # 控制生成文本的随机性
        "max_tokens": 8000  # 最大生成token数
    }
    
    try:
        response = requests.post(DEEPSEEK_API_URL, headers=headers, json=payload, timeout=60)  # 发送POST请求
        response.raise_for_status()  # 检查请求是否成功
        return response.json()["choices"][0]["message"]["content"]  # 返回API响应中的文本内容
    except Exception as e:
        raise Exception(f"DeepSeek API调用失败: {str(e)}")  # 抛出异常信息

def extract_translation_and_vocabulary(text):
    """
    使用DeepSeek API进行翻译和词汇提取
    :param text: 需要翻译的英文文本
    :return: 包含翻译结果和词汇表的字典
    """
    prompt = f"""Please perform the following tasks:
1. Translate the following English texts into Chinese accurately.
2. Extract important professional words (nouns and technical terms) from the original text to generate a glossary.
3. The output Chinese translation retains the text format of the original text.

English text: {text}

Please return the results strictly in JSON format:
{{
"translation": "Chinese translation results",
"vocabulary": [
{{
"english": "English vocabulary",
"chinese": "Chinese translation",
"explanation": "explain in detail"
}}
]
}}

Requirements:
-Translation should be accurate and fluent, and the consistency of technical terms should be maintained.
-Extract 5-10 most important professional words.
-Explain it concisely."""
    
    result = call_deepseek_api(prompt)  # 调用API获取结果
    
    try:
        # 清理API返回结果，提取JSON部分
        start_index = result.find('{')
        end_index = result.rfind('}') + 1
        json_str = result[start_index:end_index]
        return json.loads(json_str)  # 解析JSON字符串并返回
    except:
        # 如果JSON解析失败，返回默认结构
        return {
            "translation": "翻译解析失败，请重试",
            "vocabulary": []
        }

def create_word_document(original_text, translation_data, filename, include_vocabulary):
    """
    创建包含翻译结果和专业词汇表的Word文档
    :param original_text: 原始英文文本
    :param translation_data: 包含翻译结果和词汇表的字典
    :param filename: 生成的Word文档文件名
    :return: 生成的Word文档路径
    """
    doc = Document()  # 创建一个新的Word文档
    
    # 添加标题
    title = doc.add_heading('智能翻译报告', 0)  # 添加一级标题
    title.alignment = 1  # 设置标题居中
    
    # 添加基本信息表格
    doc.add_heading('基本信息', level=1)  # 添加二级标题
    info_table = doc.add_table(rows=3, cols=2)  # 创建一个3行2列的表格
    info_table.cell(0, 0).text = '生成时间'
    info_table.cell(0, 1).text = datetime.now().strftime('%Y-%m-%d %H:%M:%S')  # 当前时间
    info_table.cell(1, 0).text = '文档ID'
    info_table.cell(1, 1).text = filename.replace('.docx', '')  # 文档ID
    info_table.cell(2, 0).text = '字符数'
    info_table.cell(2, 1).text = str(len(original_text))  # 字符数统计
    
    # 添加原文内容
    doc.add_heading('原文内容', level=1)
    doc.add_paragraph(original_text)  # 添加原文段落
    
    # 添加翻译结果
    doc.add_heading('翻译结果', level=1)
    doc.add_paragraph(translation_data['translation'])  # 添加翻译结果段落
    
    # 添加词汇表
    if include_vocabulary:
        doc.add_heading('专业词汇表', level=1)
        vocab_table = doc.add_table(rows=1, cols=3)  # 创建一个1行3列的表格
        vocab_table.style = 'Light Grid Accent 1'  # 设置表格样式
        
        # 设置表头
        hdr_cells = vocab_table.rows[0].cells
        hdr_cells[0].text = '英文术语'
        hdr_cells[1].text = '中文翻译'
        hdr_cells[2].text = '解释说明'
        
        # 添加词汇行
        for vocab in translation_data['vocabulary']:
            row_cells = vocab_table.add_row().cells
            row_cells[0].text = vocab.get('english', '')
            row_cells[1].text = vocab.get('chinese', '')
            row_cells[2].text = vocab.get('explanation', '')
    
    # 保存文档
    filepath = os.path.join('downloads', filename)  # 设置文件保存路径
    doc.save(filepath)  # 保存文档
    return filepath  # 返回文件路径

@app.route('/')
def index():
    """
    主页路由
    :return: 渲染index.html模板
    """
    return render_template('index.html')  # 渲染主页模板

@app.route('/api/v1/translate', methods=['POST'])
def translate_api():
    """
    翻译API接口
    :return: JSON格式的翻译结果或错误信息
    """
    try:
        data = request.get_json()  # 获取请求的JSON数据
        
        if not data or 'text' not in data:
            return jsonify({
                "success": False,
                "error": "缺少必要的text字段"
            }), 400  # 返回400错误
        
        text = data['text'].strip()  # 获取并清理文本
        if not text:
            return jsonify({
                "success": False, 
                "error": "文本内容不能为空"
            }), 400  # 返回400错误
        
        # 调用翻译功能
        translation_data = extract_translation_and_vocabulary(text)  # 获取翻译结果
        
        # 准备响应数据
        response_data = {
            "success": True,
            "original_text": text,
            "translation": translation_data['translation'],
            "vocabulary": []
        }
        
        # 如果需要Word文档
        output_format = data.get('output_format', 'json')  # 获取输出格式
        include_vocabulary = data.get('include_vocabulary', True)  # 是否包含词汇表
        if include_vocabulary:
            response_data['vocabulary'] = translation_data.get('vocabulary', [])
        if output_format == 'word':
            filename = f"translation_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{uuid.uuid4().hex[:8]}.docx"  # 生成文件名
            filepath = create_word_document(text, translation_data, filename, include_vocabulary)  # 创建Word文档
            response_data['word_document_url'] = f"{filepath}"  # 添加下载链接
        
        return jsonify(response_data)  # 返回JSON响应
        
    except Exception as e:
        return jsonify({
            "success": False,
            "error": f"处理失败: {str(e)}"
        }), 500  # 返回500错误

@app.route('/downloads/<filename>')
def download_file(filename):
    """
    下载Word文档
    :param filename: 文件名
    :return: 文件下载响应或错误信息
    """
    try:
        return send_file(
            os.path.join('downloads', filename),  # 文件路径
            as_attachment=True,  # 作为附件下载
            download_name=filename  # 下载文件名
        )
    except FileNotFoundError:
        return jsonify({
            "success": False,
            "error": "文件未找到"
        }), 404  # 返回404错误

@app.route('/api/health')
def health_check():
    """健康检查接口"""
    return jsonify({
        "status": "healthy",
        "timestamp": datetime.now().isoformat(),
        "service": "Intelligent Translation API"
    })

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)